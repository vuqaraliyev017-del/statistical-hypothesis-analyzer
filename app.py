import io
import itertools
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import streamlit as st
from scipy import stats
import statsmodels.api as sm
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from statsmodels.stats.outliers_influence import variance_inflation_factor
from docx import Document

ALPHA_DEFAULT = 0.05

TEXTS = {
    "AZ": {
        "name": "Azərbaycan dili",
        "title": "📊 Statistik Hipotez Analiz Aləti",
        "caption": "Excel/CSV yüklə → hipotez qur → t-test, ANOVA, korrelyasiya, chi-square və regression nəticələrini avtomatik al.",
        "settings": "Ayarlar", "alpha": "Əhəmiyyətlilik səviyyəsi / alpha",
        "alpha_rule": "Qayda: p-value < alpha olduqda H0 rədd edilir / H1 qəbul edilir.",
        "upload": "Excel və ya CSV faylını yüklə", "upload_info": "Başlamaq üçün Excel və ya CSV faylı yüklə.",
        "file_error": "Fayl oxunmadı", "preview": "Data önizləmə", "rows": "Sətir", "cols": "Sütun",
        "num_cols": "Numeric sütunlar", "cat_cols": "Categorical sütunlar",
        "tabs": ["Təsviri statistika", "Avtomatik təkliflər", "t-test", "ANOVA", "Korrelyasiya", "Chi-square", "Regression", "Report export"],
        "descriptive": "Təsviri statistika", "num_vars": "Numeric dəyişənlər", "no_num": "Numeric dəyişən seçilməyib.",
        "auto_title": "Avtomatik hipotez/test təklifləri", "no_sugg": "Uyğun dəyişən kombinasiyası tapılmadı.",
        "ttype": "t-test növü", "ind": "Müstəqil qruplar", "paired": "Paired / əvvəl-sonra",
        "dep_num": "Numeric asılı dəyişən", "group2": "2 qruplu qrup dəyişəni", "run_ind": "Independent t-test çalışdır",
        "before": "Əvvəl / dəyişən 1", "after": "Sonra / dəyişən 2", "run_paired": "Paired t-test çalışdır",
        "group3": "3+ qruplu qrup dəyişəni", "run_anova": "ANOVA çalışdır", "posthoc": "Post-hoc Tukey HSD",
        "x": "Dəyişən X", "y": "Dəyişən Y", "method": "Metod", "run_corr": "Korrelyasiya çalışdır",
        "cat1": "Categorical dəyişən 1", "cat2": "Categorical dəyişən 2", "run_chi": "Chi-square çalışdır", "crosstab": "Crosstab",
        "dep_y": "Asılı dəyişən / Y", "ind_x": "Müstəqil dəyişənlər / X", "run_reg": "Regression çalışdır", "choose_x": "Ən azı bir X dəyişəni seç.",
        "coef": "Əmsallar", "full": "Tam statsmodels nəticəsi", "saved": "Yadda saxlanılmış analiz nəticələri", "no_results": "Hələ nəticə yoxdur. Əvvəl testlərdən birini çalışdır.",
        "download_xlsx": "Excel report yüklə", "download_docx": "Word report yüklə", "clear": "Nəticələri təmizlə", "cleared": "Nəticələr təmizləndi.",
        "reject": "H0 rədd edildi / H1 qəbul edildi", "not_reject": "H0 rədd edilmədi / H1 dəstəklənmədi", "no_decision": "Qərar verilə bilmədi", "not_suitable": "Uyğun deyil",
        "report_title": "Statistik Analiz Hesabatı", "report_note": "Bu report avtomatik olaraq yaradılıb. Qərarlar alpha səviyyəsinə əsasən verilib.",
        "h0_2grp": "{g1} və {g2} qrupları arasında {v} üzrə orta göstəricidə statistik fərq yoxdur.",
        "h1_2grp": "{g1} və {g2} qrupları arasında {v} üzrə orta göstəricidə statistik fərq vardır.",
        "sum_t_rej": "Independent t-test nəticəsinə görə p-value={p}. p < alpha olduğu üçün H0 rədd edilir və qruplar arasında statistik əhəmiyyətli fərq var.",
        "sum_t_not": "Independent t-test nəticəsinə görə p-value={p}. p ≥ alpha olduğu üçün H0 rədd edilmir və statistik əhəmiyyətli fərq tapılmadı.",
        "h0_pair": "{b} və {a} ölçmələri arasında orta göstəricidə statistik fərq yoxdur.", "h1_pair": "{b} və {a} ölçmələri arasında orta göstəricidə statistik fərq vardır.",
        "sum_pair_rej": "Paired t-test nəticəsinə görə p-value={p}. p < alpha olduğu üçün H0 rədd edilir və iki ölçmə arasında statistik əhəmiyyətli fərq var.",
        "sum_pair_not": "Paired t-test nəticəsinə görə p-value={p}. p ≥ alpha olduğu üçün H0 rədd edilmir və statistik əhəmiyyətli fərq tapılmadı.",
        "h0_anova": "{g} qrupları arasında {v} üzrə orta göstəricidə statistik fərq yoxdur.", "h1_anova": "{g} qruplarından ən azı biri üzrə {v} orta göstəricisi fərqlidir.",
        "sum_anova_rej": "One-way ANOVA nəticəsinə görə p-value={p}. p < alpha olduğu üçün H0 rədd edilir və qruplar arasında statistik əhəmiyyətli fərq var.",
        "sum_anova_not": "One-way ANOVA nəticəsinə görə p-value={p}. p ≥ alpha olduğu üçün H0 rədd edilmir və qruplar arasında statistik əhəmiyyətli fərq tapılmadı.",
        "h0_corr": "{x} və {y} arasında statistik əhəmiyyətli əlaqə yoxdur.", "h1_corr": "{x} və {y} arasında statistik əhəmiyyətli əlaqə vardır.",
        "sum_corr": "{test} nəticəsinə görə correlation coefficient={r}, p-value={p}. {sig} Əlaqənin istiqaməti {direction}, gücü isə {strength} səviyyədədir.",
        "sig_yes": "p < alpha olduğu üçün statistik əhəmiyyətli əlaqə var.", "sig_no": "p ≥ alpha olduğu üçün statistik əhəmiyyətli əlaqə tapılmadı.",
        "pos": "müsbət", "neg": "mənfi", "weak": "zəif", "medium": "orta", "strong": "güclü",
        "h0_chi": "{c1} və {c2} dəyişənləri arasında asılılıq yoxdur.", "h1_chi": "{c1} və {c2} dəyişənləri arasında statistik əhəmiyyətli asılılıq vardır.",
        "sum_chi_rej": "Chi-square test nəticəsinə görə p-value={p}. p < alpha olduğu üçün H0 rədd edilir və dəyişənlər arasında statistik əhəmiyyətli asılılıq var.",
        "sum_chi_not": "Chi-square test nəticəsinə görə p-value={p}. p ≥ alpha olduğu üçün H0 rədd edilmir və statistik əhəmiyyətli asılılıq tapılmadı.",
        "h0_reg": "{xs} dəyişənləri birlikdə {y} dəyişənini statistik əhəmiyyətli izah etmir.", "h1_reg": "{xs} dəyişənlərindən ən azı biri {y} dəyişənini statistik əhəmiyyətli izah edir.",
        "sum_reg_sig": "Linear regression modelinin F-test p-value={p}, R²={r2}, adjusted R²={ar2}. Model ümumi olaraq statistik əhəmiyyətlidir.",
        "sum_reg_not": "Linear regression modelinin F-test p-value={p}, R²={r2}, adjusted R²={ar2}. Model ümumi olaraq statistik əhəmiyyətli deyil.",
        "need2": "{g} sütununda tam 2 qrup olmalıdır. Hazırda {n} qrup var.", "need3": "ANOVA üçün ən azı 3 qrup lazımdır. 2 qrup varsa t-test istifadə edin.",
        "need_obs": "Hər qrupda ən azı 2 müşahidə olmalıdır.", "need_pair": "Ən azı 2 paired müşahidə lazımdır.", "need_corr": "Korrelyasiya üçün ən azı 3 müşahidə lazımdır.",
        "need_cat": "Hər iki categorical dəyişəndə ən azı 2 kateqoriya olmalıdır.", "no_data": "Məlumat yoxdur.", "too_few": "Müşahidə sayı çox azdır.",
        "assumption": "Assumption check", "normality_few": "Normality üçün müşahidə sayı azdır"
    },
    "TR": {},
    "EN": {}
}

TEXTS["TR"] = {**TEXTS["AZ"], **{
    "name": "Türkçe", "title": "📊 İstatistiksel Hipotez Analiz Aracı",
    "caption": "Excel/CSV yükle → hipotez oluştur → t-test, ANOVA, korelasyon, chi-square ve regresyon sonuçlarını otomatik al.",
    "settings": "Ayarlar", "alpha": "Anlamlılık düzeyi / alpha", "alpha_rule": "Kural: p-value < alpha olduğunda H0 reddedilir / H1 kabul edilir.",
    "upload": "Excel veya CSV dosyasını yükle", "upload_info": "Başlamak için Excel veya CSV dosyası yükle.", "file_error": "Dosya okunamadı", "preview": "Veri önizleme", "rows": "Satır", "cols": "Sütun",
    "num_cols": "Sayısal sütunlar", "cat_cols": "Kategorik sütunlar", "tabs": ["Tanımlayıcı istatistikler", "Otomatik öneriler", "t-test", "ANOVA", "Korelasyon", "Chi-square", "Regresyon", "Rapor çıktısı"],
    "descriptive": "Tanımlayıcı istatistikler", "num_vars": "Sayısal değişkenler", "no_num": "Sayısal değişken seçilmedi.", "auto_title": "Otomatik hipotez/test önerileri", "no_sugg": "Uygun değişken kombinasyonu bulunamadı.",
    "ttype": "t-test türü", "ind": "Bağımsız gruplar", "paired": "Eşleştirilmiş / önce-sonra", "dep_num": "Sayısal bağımlı değişken", "group2": "2 gruplu grup değişkeni", "run_ind": "Independent t-test çalıştır",
    "before": "Önce / değişken 1", "after": "Sonra / değişken 2", "run_paired": "Paired t-test çalıştır", "group3": "3+ gruplu grup değişkeni", "run_anova": "ANOVA çalıştır",
    "x": "Değişken X", "y": "Değişken Y", "run_corr": "Korelasyon çalıştır", "cat1": "Kategorik değişken 1", "cat2": "Kategorik değişken 2", "run_chi": "Chi-square çalıştır", "crosstab": "Çapraz tablo",
    "dep_y": "Bağımlı değişken / Y", "ind_x": "Bağımsız değişkenler / X", "run_reg": "Regresyon çalıştır", "choose_x": "En az bir X değişkeni seç.", "coef": "Katsayılar", "saved": "Kaydedilmiş analiz sonuçları",
    "no_results": "Henüz sonuç yok. Önce testlerden birini çalıştır.", "download_xlsx": "Excel raporu indir", "download_docx": "Word raporu indir", "clear": "Sonuçları temizle", "cleared": "Sonuçlar temizlendi.",
    "reject": "H0 reddedildi / H1 kabul edildi", "not_reject": "H0 reddedilemedi / H1 desteklenmedi", "no_decision": "Karar verilemedi", "not_suitable": "Uygun değil",
    "report_title": "İstatistiksel Analiz Raporu", "report_note": "Bu rapor otomatik olarak oluşturulmuştur. Kararlar alpha düzeyine göre verilmiştir.",
    "h0_2grp": "{g1} ve {g2} grupları arasında {v} ortalaması açısından istatistiksel fark yoktur.", "h1_2grp": "{g1} ve {g2} grupları arasında {v} ortalaması açısından istatistiksel fark vardır.",
    "sum_t_rej": "Independent t-test sonucuna göre p-value={p}. p < alpha olduğu için H0 reddedilir ve gruplar arasında istatistiksel olarak anlamlı fark vardır.",
    "sum_t_not": "Independent t-test sonucuna göre p-value={p}. p ≥ alpha olduğu için H0 reddedilemez ve istatistiksel olarak anlamlı fark bulunmamıştır.",
    "h0_pair": "{b} ve {a} ölçümleri arasında ortalama açısından istatistiksel fark yoktur.", "h1_pair": "{b} ve {a} ölçümleri arasında ortalama açısından istatistiksel fark vardır.",
    "sum_pair_rej": "Paired t-test sonucuna göre p-value={p}. p < alpha olduğu için H0 reddedilir ve iki ölçüm arasında istatistiksel olarak anlamlı fark vardır.",
    "sum_pair_not": "Paired t-test sonucuna göre p-value={p}. p ≥ alpha olduğu için H0 reddedilemez ve istatistiksel olarak anlamlı fark bulunmamıştır.",
    "h0_anova": "{g} grupları arasında {v} ortalaması açısından istatistiksel fark yoktur.", "h1_anova": "{g} gruplarından en az birinde {v} ortalaması farklıdır.",
    "sum_anova_rej": "One-way ANOVA sonucuna göre p-value={p}. p < alpha olduğu için H0 reddedilir ve gruplar arasında istatistiksel olarak anlamlı fark vardır.",
    "sum_anova_not": "One-way ANOVA sonucuna göre p-value={p}. p ≥ alpha olduğu için H0 reddedilemez ve gruplar arasında istatistiksel olarak anlamlı fark bulunmamıştır.",
    "h0_corr": "{x} ve {y} arasında istatistiksel olarak anlamlı ilişki yoktur.", "h1_corr": "{x} ve {y} arasında istatistiksel olarak anlamlı ilişki vardır.",
    "sum_corr": "{test} sonucuna göre korelasyon katsayısı={r}, p-value={p}. {sig} İlişkinin yönü {direction}, gücü ise {strength} düzeydedir.",
    "sig_yes": "p < alpha olduğu için istatistiksel olarak anlamlı ilişki vardır.", "sig_no": "p ≥ alpha olduğu için istatistiksel olarak anlamlı ilişki bulunmamıştır.",
    "pos": "pozitif", "neg": "negatif", "weak": "zayıf", "medium": "orta", "strong": "güçlü",
    "h0_chi": "{c1} ve {c2} değişkenleri arasında ilişki yoktur.", "h1_chi": "{c1} ve {c2} değişkenleri arasında istatistiksel olarak anlamlı ilişki vardır.",
    "sum_chi_rej": "Chi-square test sonucuna göre p-value={p}. p < alpha olduğu için H0 reddedilir ve değişkenler arasında istatistiksel olarak anlamlı ilişki vardır.",
    "sum_chi_not": "Chi-square test sonucuna göre p-value={p}. p ≥ alpha olduğu için H0 reddedilemez ve istatistiksel olarak anlamlı ilişki bulunmamıştır.",
    "h0_reg": "{xs} değişkenleri birlikte {y} değişkenini istatistiksel olarak anlamlı biçimde açıklamamaktadır.", "h1_reg": "{xs} değişkenlerinden en az biri {y} değişkenini istatistiksel olarak anlamlı biçimde açıklamaktadır.",
    "sum_reg_sig": "Linear regression modelinin F-test p-value={p}, R²={r2}, adjusted R²={ar2}. Model genel olarak istatistiksel olarak anlamlıdır.",
    "sum_reg_not": "Linear regression modelinin F-test p-value={p}, R²={r2}, adjusted R²={ar2}. Model genel olarak istatistiksel olarak anlamlı değildir.",
    "need2": "{g} sütununda tam 2 grup olmalıdır. Mevcut grup sayısı: {n}.", "need3": "ANOVA için en az 3 grup gereklidir. 2 grup varsa t-test kullanın.",
    "need_obs": "Her grupta en az 2 gözlem olmalıdır.", "need_pair": "En az 2 eşleştirilmiş gözlem gereklidir.", "need_corr": "Korelasyon için en az 3 gözlem gereklidir.",
    "need_cat": "Her iki kategorik değişkende de en az 2 kategori olmalıdır.", "no_data": "Veri yok.", "too_few": "Gözlem sayısı çok azdır.", "assumption": "Varsayım kontrolü", "normality_few": "Normallik için gözlem sayısı azdır"
}}

TEXTS["EN"] = {**TEXTS["AZ"], **{
    "name": "English", "title": "📊 Statistical Hypothesis Analyzer", "caption": "Upload Excel/CSV → generate hypotheses → automatically run t-test, ANOVA, correlation, chi-square, and regression.",
    "settings": "Settings", "alpha": "Significance level / alpha", "alpha_rule": "Rule: when p-value < alpha, H0 is rejected / H1 is accepted.",
    "upload": "Upload Excel or CSV file", "upload_info": "Upload an Excel or CSV file to start.", "file_error": "File could not be read", "preview": "Data preview", "rows": "Rows", "cols": "Columns",
    "num_cols": "Numeric columns", "cat_cols": "Categorical columns", "tabs": ["Descriptive statistics", "Auto suggestions", "t-test", "ANOVA", "Correlation", "Chi-square", "Regression", "Report export"],
    "descriptive": "Descriptive statistics", "num_vars": "Numeric variables", "no_num": "No numeric variable selected.", "auto_title": "Automatic hypothesis/test suggestions", "no_sugg": "No suitable variable combination was found.",
    "ttype": "t-test type", "ind": "Independent samples", "paired": "Paired / before-after", "dep_num": "Numeric dependent variable", "group2": "Grouping variable with 2 groups", "run_ind": "Run independent t-test",
    "before": "Before / variable 1", "after": "After / variable 2", "run_paired": "Run paired t-test", "group3": "Grouping variable with 3+ groups", "run_anova": "Run ANOVA",
    "x": "Variable X", "y": "Variable Y", "run_corr": "Run correlation", "cat1": "Categorical variable 1", "cat2": "Categorical variable 2", "run_chi": "Run chi-square", "crosstab": "Crosstab",
    "dep_y": "Dependent variable / Y", "ind_x": "Independent variables / X", "run_reg": "Run regression", "choose_x": "Select at least one X variable.", "coef": "Coefficients", "saved": "Saved analysis results",
    "no_results": "No results yet. Run one of the tests first.", "download_xlsx": "Download Excel report", "download_docx": "Download Word report", "clear": "Clear results", "cleared": "Results cleared.",
    "reject": "H0 rejected / H1 accepted", "not_reject": "H0 not rejected / H1 not supported", "no_decision": "Decision unavailable", "not_suitable": "Not suitable",
    "report_title": "Statistical Analysis Report", "report_note": "This report was generated automatically. Decisions are based on the selected alpha level.",
    "h0_2grp": "There is no statistically significant difference in the mean of {v} between {g1} and {g2}.", "h1_2grp": "There is a statistically significant difference in the mean of {v} between {g1} and {g2}.",
    "sum_t_rej": "According to the independent t-test, p-value={p}. Since p < alpha, H0 is rejected and there is a statistically significant difference between the groups.",
    "sum_t_not": "According to the independent t-test, p-value={p}. Since p ≥ alpha, H0 is not rejected and no statistically significant difference was found.",
    "h0_pair": "There is no statistically significant mean difference between {b} and {a}.", "h1_pair": "There is a statistically significant mean difference between {b} and {a}.",
    "sum_pair_rej": "According to the paired t-test, p-value={p}. Since p < alpha, H0 is rejected and there is a statistically significant difference between the two measurements.",
    "sum_pair_not": "According to the paired t-test, p-value={p}. Since p ≥ alpha, H0 is not rejected and no statistically significant difference was found.",
    "h0_anova": "There is no statistically significant difference in the mean of {v} across {g} groups.", "h1_anova": "At least one {g} group has a different mean of {v}.",
    "sum_anova_rej": "According to the one-way ANOVA, p-value={p}. Since p < alpha, H0 is rejected and there is a statistically significant difference between groups.",
    "sum_anova_not": "According to the one-way ANOVA, p-value={p}. Since p ≥ alpha, H0 is not rejected and no statistically significant difference was found between groups.",
    "h0_corr": "There is no statistically significant relationship between {x} and {y}.", "h1_corr": "There is a statistically significant relationship between {x} and {y}.",
    "sum_corr": "According to the {test}, correlation coefficient={r}, p-value={p}. {sig} The relationship is {direction} and its strength is {strength}.",
    "sig_yes": "Since p < alpha, there is a statistically significant relationship.", "sig_no": "Since p ≥ alpha, no statistically significant relationship was found.",
    "pos": "positive", "neg": "negative", "weak": "weak", "medium": "moderate", "strong": "strong",
    "h0_chi": "There is no association between {c1} and {c2}.", "h1_chi": "There is a statistically significant association between {c1} and {c2}.",
    "sum_chi_rej": "According to the chi-square test, p-value={p}. Since p < alpha, H0 is rejected and there is a statistically significant association between the variables.",
    "sum_chi_not": "According to the chi-square test, p-value={p}. Since p ≥ alpha, H0 is not rejected and no statistically significant association was found.",
    "h0_reg": "The variables {xs} do not jointly explain {y} in a statistically significant way.", "h1_reg": "At least one of the variables {xs} explains {y} in a statistically significant way.",
    "sum_reg_sig": "The linear regression model has F-test p-value={p}, R²={r2}, adjusted R²={ar2}. Overall, the model is statistically significant.",
    "sum_reg_not": "The linear regression model has F-test p-value={p}, R²={r2}, adjusted R²={ar2}. Overall, the model is not statistically significant.",
    "need2": "The {g} column must contain exactly 2 groups. Current number of groups: {n}.", "need3": "ANOVA requires at least 3 groups. If there are 2 groups, use a t-test.",
    "need_obs": "Each group must have at least 2 observations.", "need_pair": "At least 2 paired observations are required.", "need_corr": "Correlation requires at least 3 observations.",
    "need_cat": "Both categorical variables must have at least 2 categories.", "no_data": "No data available.", "too_few": "The number of observations is too small.", "assumption": "Assumption check", "normality_few": "Too few observations for normality check"
}}

@dataclass
class AnalysisResult:
    analysis_type: str
    variables: str
    h0: str
    h1: str
    statistic_name: str
    statistic_value: Optional[float]
    p_value: Optional[float]
    decision: str
    summary: str
    details: str = ""

def fp(p):
    if p is None or (isinstance(p, float) and np.isnan(p)): return "N/A"
    return "< 0.001" if p < 0.001 else f"{p:.4f}"

def decision(p, alpha, t):
    if p is None or (isinstance(p, float) and np.isnan(p)): return t["no_decision"]
    return t["reject"] if p < alpha else t["not_reject"]

def load_data(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".csv"): df = pd.read_csv(uploaded_file)
    elif name.endswith((".xlsx", ".xls")): df = pd.read_excel(uploaded_file)
    else: raise ValueError("Only CSV and Excel files are supported.")
    df = df.dropna(how="all")
    df.columns = [str(c).strip() for c in df.columns]
    return df

def detect_columns(df):
    nums = df.select_dtypes(include=[np.number]).columns.tolist()
    cats = [c for c in df.columns if c not in nums]
    for c in df.columns:
        if c not in nums:
            conv = pd.to_numeric(df[c], errors="coerce")
            if conv.notna().mean() >= 0.75:
                df[c] = conv; nums.append(c)
                if c in cats: cats.remove(c)
    for c in nums.copy():
        if df[c].nunique(dropna=True) <= 10 and c not in cats: cats.append(c)
    return nums, cats

def normality(df, num, t, group=None):
    out = []
    if group:
        for g, sub in df[[num, group]].dropna().groupby(group):
            vals = sub[num]
            out.append(f"{group}={g}: Shapiro-Wilk p={fp(stats.shapiro(vals)[1])}" if len(vals) >= 3 else f"{group}={g}: {t['normality_few']}")
    else:
        vals = df[num].dropna()
        out.append(f"Shapiro-Wilk p={fp(stats.shapiro(vals)[1])}" if len(vals) >= 3 else t["normality_few"])
    return "; ".join(out)

def desc_stats(df, nums):
    if not nums: return pd.DataFrame()
    d = df[nums].describe().T
    d["missing"] = df[nums].isna().sum(); d["median"] = df[nums].median(); d["skewness"] = df[nums].skew(); d["kurtosis"] = df[nums].kurtosis()
    return d.reset_index().rename(columns={"index":"variable"})

def ind_ttest(df, num, group, alpha, t):
    data = df[[num, group]].dropna(); groups = list(data[group].dropna().unique())
    if len(groups) != 2: return AnalysisResult("Independent Samples t-test", f"{num} by {group}", "", "", "t", None, None, t["not_suitable"], t["need2"].format(g=group, n=len(groups)))
    g1, g2 = groups[0], groups[1]; x1 = data[data[group]==g1][num]; x2 = data[data[group]==g2][num]
    if len(x1)<2 or len(x2)<2: return AnalysisResult("Independent Samples t-test", f"{num} by {group}", "", "", "t", None, None, t["not_suitable"], t["need_obs"])
    stat, p = stats.ttest_ind(x1, x2, equal_var=False, nan_policy="omit")
    return AnalysisResult("Independent Samples t-test", f"{num} by {group}", t["h0_2grp"].format(g1=g1,g2=g2,v=num), t["h1_2grp"].format(g1=g1,g2=g2,v=num), "t", float(stat), float(p), decision(float(p), alpha, t), (t["sum_t_rej"] if p<alpha else t["sum_t_not"]).format(p=fp(p)), f"{g1}: n={len(x1)}, mean={x1.mean():.4f}, std={x1.std():.4f}; {g2}: n={len(x2)}, mean={x2.mean():.4f}, std={x2.std():.4f}. {t['assumption']}: {normality(data,num,t,group)}")

def paired_ttest(df, before, after, alpha, t):
    data = df[[before, after]].dropna()
    if len(data)<2: return AnalysisResult("Paired Samples t-test", f"{before} vs {after}", "", "", "t", None, None, t["not_suitable"], t["need_pair"])
    stat,p = stats.ttest_rel(data[before], data[after]); diff = data[after]-data[before]
    return AnalysisResult("Paired Samples t-test", f"{before} vs {after}", t["h0_pair"].format(b=before,a=after), t["h1_pair"].format(b=before,a=after), "t", float(stat), float(p), decision(float(p), alpha, t), (t["sum_pair_rej"] if p<alpha else t["sum_pair_not"]).format(p=fp(p)), f"n={len(data)}, mean difference={diff.mean():.4f}, std difference={diff.std():.4f}. {t['assumption']}: {normality(pd.DataFrame({'diff':diff}), 'diff', t)}")

def anova(df, num, group, alpha, t):
    data = df[[num, group]].dropna(); grouped = [s[num].values for _, s in data.groupby(group)]
    if len(grouped)<3: return AnalysisResult("One-way ANOVA", f"{num} by {group}", "", "", "F", None, None, t["not_suitable"], t["need3"]), None
    if any(len(x)<2 for x in grouped): return AnalysisResult("One-way ANOVA", f"{num} by {group}", "", "", "F", None, None, t["not_suitable"], t["need_obs"]), None
    stat,p = stats.f_oneway(*grouped); tuk = None
    if p < alpha:
        try:
            res = pairwise_tukeyhsd(endog=data[num], groups=data[group], alpha=alpha)
            tuk = pd.DataFrame(data=res.summary().data[1:], columns=res.summary().data[0])
        except Exception: tuk = None
    return AnalysisResult("One-way ANOVA", f"{num} by {group}", t["h0_anova"].format(g=group,v=num), t["h1_anova"].format(g=group,v=num), "F", float(stat), float(p), decision(float(p), alpha, t), (t["sum_anova_rej"] if p<alpha else t["sum_anova_not"]).format(p=fp(p)), f"{t['assumption']}: {normality(data,num,t,group)}"), tuk

def corr(df, x, y, alpha, method, t):
    data = df[[x,y]].dropna()
    if len(data)<3: return AnalysisResult("Correlation", f"{x} and {y}", "", "", "r", None, None, t["not_suitable"], t["need_corr"])
    if method == "spearman": stat,p = stats.spearmanr(data[x], data[y]); test="Spearman correlation"; sname="rho"
    else: stat,p = stats.pearsonr(data[x], data[y]); test="Pearson correlation"; sname="r"
    direction = t["pos"] if stat>0 else t["neg"]; strength = t["weak"] if abs(stat)<0.3 else t["medium"] if abs(stat)<0.7 else t["strong"]
    return AnalysisResult(test, f"{x} and {y}", t["h0_corr"].format(x=x,y=y), t["h1_corr"].format(x=x,y=y), sname, float(stat), float(p), decision(float(p), alpha, t), t["sum_corr"].format(test=test,r=f"{stat:.4f}",p=fp(p),sig=t["sig_yes"] if p<alpha else t["sig_no"],direction=direction,strength=strength), f"n={len(data)}")

def chi(df, c1, c2, alpha, t):
    data = df[[c1,c2]].dropna()
    if data.empty: return AnalysisResult("Chi-square test", f"{c1} and {c2}", "", "", "chi2", None, None, t["not_suitable"], t["no_data"]), None
    table = pd.crosstab(data[c1], data[c2])
    if table.shape[0]<2 or table.shape[1]<2: return AnalysisResult("Chi-square test", f"{c1} and {c2}", "", "", "chi2", None, None, t["not_suitable"], t["need_cat"]), table
    stat,p,dof,_ = stats.chi2_contingency(table)
    return AnalysisResult("Chi-square test", f"{c1} and {c2}", t["h0_chi"].format(c1=c1,c2=c2), t["h1_chi"].format(c1=c1,c2=c2), "chi2", float(stat), float(p), decision(float(p), alpha, t), (t["sum_chi_rej"] if p<alpha else t["sum_chi_not"]).format(p=fp(p)), f"chi-square={stat:.4f}, dof={dof}, n={len(data)}"), table

def regression(df, y, xs, alpha, t):
    data = df[[y]+xs].dropna()
    if len(data) <= len(xs)+2: return AnalysisResult("Linear Regression", f"{y} ~ {', '.join(xs)}", "", "", "F", None, None, t["not_suitable"], t["too_few"]), pd.DataFrame(), ""
    X = sm.add_constant(data[xs]); model = sm.OLS(data[y], X).fit(); p = float(model.f_pvalue) if model.f_pvalue is not None else np.nan; xst = ", ".join(xs)
    coef = pd.DataFrame({"variable": model.params.index, "coefficient": model.params.values, "p_value": model.pvalues.values, "significant": ["Yes" if v<alpha else "No" for v in model.pvalues.values]})
    details = " ".join([f"{r.variable}: coefficient={r.coefficient:.4f}, p={fp(r.p_value)}." for r in coef.itertuples() if r.variable != "const"])
    try:
        if len(xs)>=2: details += " VIF: " + "; ".join([f"{n}={variance_inflation_factor(data[xs].values,i):.2f}" for i,n in enumerate(xs)])
    except Exception: pass
    return AnalysisResult("Linear Regression", f"{y} ~ {xst}", t["h0_reg"].format(xs=xst,y=y), t["h1_reg"].format(xs=xst,y=y), "F", float(model.fvalue), p, decision(p, alpha, t), (t["sum_reg_sig"] if p<alpha else t["sum_reg_not"]).format(p=fp(p),r2=f"{model.rsquared:.4f}",ar2=f"{model.rsquared_adj:.4f}"), details), coef, model.summary().as_text()

def suggestions(df, nums, cats):
    out=[]
    for g in cats:
        n=df[g].nunique(dropna=True)
        for num in nums:
            if g!=num and n==2: out.append(f"Independent t-test: {num} by {g}")
            elif g!=num and 3<=n<=10: out.append(f"One-way ANOVA: {num} by {g}")
    for x,y in itertools.combinations(nums,2): out.append(f"Correlation: {x} and {y}")
    for c1,c2 in itertools.combinations(cats,2):
        if df[c1].nunique(dropna=True)<=20 and df[c2].nunique(dropna=True)<=20: out.append(f"Chi-square: {c1} and {c2}")
    return out[:100]

def to_dict(r):
    return {"Analysis Type":r.analysis_type,"Variables":r.variables,"H0":r.h0,"H1":r.h1,"Statistic":r.statistic_name,"Statistic Value":r.statistic_value,"p-value":r.p_value,"Decision":r.decision,"Summary":r.summary,"Details":r.details}

def excel_report(results, desc, extra):
    out=io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as w:
        if desc is not None and not desc.empty: desc.to_excel(w, sheet_name="Descriptive", index=False)
        if results: pd.DataFrame([to_dict(r) for r in results]).to_excel(w, sheet_name="Hypothesis Tests", index=False)
        for name, tab in (extra or {}).items(): tab.to_excel(w, sheet_name=name[:31].replace("/","_"), index=True)
    out.seek(0); return out.getvalue()

def word_report(results, desc, t):
    doc=Document(); doc.add_heading(t["report_title"],0); doc.add_paragraph(t["report_note"])
    if desc is not None and not desc.empty:
        doc.add_heading(t["descriptive"], level=1); table=doc.add_table(rows=1, cols=len(desc.columns))
        for i,c in enumerate(desc.columns): table.rows[0].cells[i].text=str(c)
        for _,row in desc.head(30).iterrows():
            cells=table.add_row().cells
            for i,c in enumerate(desc.columns):
                val=row[c]; cells[i].text=f"{val:.4f}" if isinstance(val,(float,np.floating)) else str(val)
    if results:
        doc.add_heading("Hypothesis Test Results", level=1)
        for i,r in enumerate(results,1):
            doc.add_heading(f"{i}. {r.analysis_type}: {r.variables}", level=2)
            for label,val in [("H0",r.h0),("H1",r.h1),(r.statistic_name,r.statistic_value),("p-value",fp(r.p_value)),("Decision",r.decision),("Summary",r.summary),("Details",r.details)]:
                if val not in [None,""]: doc.add_paragraph(f"{label}: {val}")
    out=io.BytesIO(); doc.save(out); out.seek(0); return out.getvalue()

st.set_page_config(page_title="Statistical Hypothesis Analyzer", layout="wide")
with st.sidebar:
    lang = st.selectbox("Language / Dil", ["AZ","TR","EN"], format_func=lambda k: f"{k} - {TEXTS[k]['name']}")
    t=TEXTS[lang]
    st.header(t["settings"]); alpha=st.number_input(t["alpha"], min_value=0.001, max_value=0.20, value=ALPHA_DEFAULT, step=0.01); st.info(t["alpha_rule"])

st.title(t["title"]); st.caption(t["caption"])
file = st.file_uploader(t["upload"], type=["csv","xlsx","xls"])
if not file:
    st.info(t["upload_info"]); st.stop()
try: df=load_data(file)
except Exception as e: st.error(f"{t['file_error']}: {e}"); st.stop()
nums,cats=detect_columns(df)
st.subheader(t["preview"]); st.write(f"{t['rows']}: **{df.shape[0]}**, {t['cols']}: **{df.shape[1]}**"); st.dataframe(df.head(100), use_container_width=True)
c1,c2=st.columns(2)
with c1: st.markdown(f"**{t['num_cols']}**"); st.write(nums)
with c2: st.markdown(f"**{t['cat_cols']}**"); st.write(cats)

tabs=st.tabs(t["tabs"])
with tabs[0]:
    st.subheader(t["descriptive"]); selected=st.multiselect(t["num_vars"], nums, default=nums[:5]); d=desc_stats(df, selected)
    st.dataframe(d, use_container_width=True) if not d.empty else st.warning(t["no_num"])
with tabs[1]:
    st.subheader(t["auto_title"]); s=suggestions(df,nums,cats)
    [st.write("• "+x) for x in s] if s else st.warning(t["no_sugg"])
with tabs[2]:
    st.subheader(t["tabs"][2]); typ=st.radio(t["ttype"], [t["ind"],t["paired"]], horizontal=True)
    if typ==t["ind"]:
        a,b=st.columns(2); num=a.selectbox(t["dep_num"], nums, key="ttn"); group=b.selectbox(t["group2"], cats, key="ttg")
        if st.button(t["run_ind"]): r=ind_ttest(df,num,group,alpha,t); st.session_state.setdefault("results",[]).append(r); st.success(r.decision); st.write(to_dict(r))
    else:
        a,b=st.columns(2); before=a.selectbox(t["before"], nums, key="pb"); after=b.selectbox(t["after"], nums, key="pa")
        if st.button(t["run_paired"]): r=paired_ttest(df,before,after,alpha,t); st.session_state.setdefault("results",[]).append(r); st.success(r.decision); st.write(to_dict(r))
with tabs[3]:
    st.subheader(t["tabs"][3]); a,b=st.columns(2); num=a.selectbox(t["dep_num"], nums, key="an"); group=b.selectbox(t["group3"], cats, key="ag")
    if st.button(t["run_anova"]):
        r,tuk=anova(df,num,group,alpha,t); st.session_state.setdefault("results",[]).append(r); st.success(r.decision); st.write(to_dict(r))
        if tuk is not None: st.markdown(f"**{t['posthoc']}**"); st.dataframe(tuk, use_container_width=True); st.session_state.setdefault("extra",{})[f"Tukey_{num}_{group}"]=tuk
with tabs[4]:
    st.subheader(t["tabs"][4]); a,b,c=st.columns(3); x=a.selectbox(t["x"], nums, key="cx"); y=b.selectbox(t["y"], nums, key="cy"); method=c.selectbox(t["method"], ["pearson","spearman"])
    if st.button(t["run_corr"]): r=corr(df,x,y,alpha,method,t); st.session_state.setdefault("results",[]).append(r); st.success(r.decision); st.write(to_dict(r))
with tabs[5]:
    st.subheader(t["tabs"][5]); a,b=st.columns(2); ca=a.selectbox(t["cat1"], cats, key="ch1"); cb=b.selectbox(t["cat2"], cats, key="ch2")
    if st.button(t["run_chi"]):
        r,tab=chi(df,ca,cb,alpha,t); st.session_state.setdefault("results",[]).append(r); st.success(r.decision); st.write(to_dict(r))
        if tab is not None: st.markdown(f"**{t['crosstab']}**"); st.dataframe(tab, use_container_width=True); st.session_state.setdefault("extra",{})[f"Crosstab_{ca}_{cb}"]=tab
with tabs[6]:
    st.subheader(t["tabs"][6]); a,b=st.columns(2); y=a.selectbox(t["dep_y"], nums, key="ry"); xs=b.multiselect(t["ind_x"], [n for n in nums if n!=y], key="rx")
    if st.button(t["run_reg"]):
        if not xs: st.warning(t["choose_x"])
        else:
            r,coef,mt=regression(df,y,xs,alpha,t); st.session_state.setdefault("results",[]).append(r); st.success(r.decision); st.write(to_dict(r))
            if not coef.empty: st.markdown(f"**{t['coef']}**"); st.dataframe(coef, use_container_width=True); st.session_state.setdefault("extra",{})[f"Regression_{y}"]=coef
            with st.expander(t["full"]): st.text(mt)
with tabs[7]:
    st.subheader(t["tabs"][7]); results=st.session_state.get("results",[]); extra=st.session_state.get("extra",{}); d=desc_stats(df, nums)
    if results: st.markdown(f"**{t['saved']}**"); st.dataframe(pd.DataFrame([to_dict(r) for r in results]), use_container_width=True)
    else: st.info(t["no_results"])
    a,b=st.columns(2)
    with a: st.download_button(t["download_xlsx"], data=excel_report(results,d,extra), file_name="statistical_analysis_report.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with b: st.download_button(t["download_docx"], data=word_report(results,d,t), file_name="statistical_analysis_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if st.button(t["clear"]): st.session_state["results"]=[]; st.session_state["extra"]={}; st.success(t["cleared"])
